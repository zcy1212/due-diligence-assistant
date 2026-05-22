import os
import zipfile
import io
from pathlib import Path
from typing import List, Dict, Tuple, Optional

try:
    from PyPDF2 import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False


class FileReader:
    def __init__(self, input_dir: str, temp_dir: str):
        self.input_dir = Path(input_dir)
        self.temp_dir = Path(temp_dir)
        self.temp_dir.mkdir(exist_ok=True)

    def extract_zip(self, zip_path: Path) -> Path:
        extract_dir = self.temp_dir / zip_path.stem
        extract_dir.mkdir(exist_ok=True)
        
        with zipfile.ZipFile(zip_path, 'r') as zip_ref:
            zip_ref.extractall(extract_dir)
        
        return extract_dir

    def find_all_files(self, directory: Path) -> List[Path]:
        files = []
        for root, _, filenames in os.walk(directory):
            for filename in filenames:
                if filename.lower().endswith(('.pdf', '.docx', '.doc', '.txt')):
                    files.append(Path(root) / filename)
        return sorted(files)

    def read_pdf(self, file_path: Path) -> str:
        text = ""
        
        if PYMUPDF_AVAILABLE:
            try:
                doc = fitz.open(file_path)
                for page in doc:
                    text += page.get_text()
                return text
            except Exception as e:
                pass
        
        if PDF_AVAILABLE:
            try:
                reader = PdfReader(file_path)
                for page in reader.pages:
                    text += page.extract_text() or ""
                return text
            except Exception as e:
                pass
        
        return f"[无法读取PDF文件: {file_path.name}]"

    def read_docx(self, file_path: Path) -> str:
        if not DOCX_AVAILABLE:
            return f"[无法读取Word文件: {file_path.name}]"
        
        try:
            doc = Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            
            for table in doc.tables:
                text += "\n[表格]\n"
                for row in table.rows:
                    text += " | ".join([cell.text for cell in row.cells]) + "\n"
            
            return text
        except Exception as e:
            return f"[读取Word文件出错: {file_path.name}]"

    def read_txt(self, file_path: Path) -> str:
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            try:
                with open(file_path, 'r', encoding='gbk') as f:
                    return f.read()
            except Exception as e:
                return f"[读取文本文件出错: {file_path.name}]"

    def read_file(self, file_path: Path) -> Dict:
        ext = file_path.suffix.lower()
        
        if ext == '.pdf':
            content = self.read_pdf(file_path)
        elif ext == '.docx':
            content = self.read_docx(file_path)
        elif ext == '.txt':
            content = self.read_txt(file_path)
        else:
            content = f"[不支持的文件格式: {ext}]"
        
        return {
            'path': str(file_path),
            'name': file_path.name,
            'type': ext,
            'content': content,
            'size': file_path.stat().st_size
        }

    def scan_directory(self, folder_name: Optional[str] = None) -> Tuple[List[Dict], Path]:
        target_dir = self.input_dir
        
        if folder_name:
            target_dir = self.input_dir / folder_name
        
        if not target_dir.exists():
            zip_files = list(self.input_dir.glob("*.zip"))
            if zip_files:
                target_dir = self.extract_zip(zip_files[0])
            else:
                raise FileNotFoundError(f"未找到目录或压缩包: {folder_name}")
        
        files = self.find_all_files(target_dir)
        
        results = []
        for file_path in files:
            results.append(self.read_file(file_path))
        
        return results, target_dir

    def get_file_list(self, folder_name: Optional[str] = None) -> List[Dict]:
        files, _ = self.scan_directory(folder_name)
        return [{'name': f['name'], 'path': f['path'], 'type': f['type'], 'size': f['size']} for f in files]

    def cleanup_temp(self):
        import shutil
        if self.temp_dir.exists():
            shutil.rmtree(self.temp_dir)
        self.temp_dir.mkdir(exist_ok=True)


if __name__ == "__main__":
    import sys
    import json
    
    input_dir = sys.argv[1] if len(sys.argv) > 1 else "input"
    temp_dir = sys.argv[2] if len(sys.argv) > 2 else "temp"
    folder_name = sys.argv[3] if len(sys.argv) > 3 else None
    action = sys.argv[4] if len(sys.argv) > 4 else "scan"
    
    reader = FileReader(input_dir, temp_dir)
    
    if action == "list":
        files = reader.get_file_list(folder_name)
        print(json.dumps(files, ensure_ascii=False, indent=2))
    elif action == "scan":
        files, _ = reader.scan_directory(folder_name)
        output = []
        for f in files:
            output.append({
                'name': f['name'],
                'path': f['path'],
                'type': f['type'],
                'content_preview': f['content'][:1000] if len(f['content']) > 1000 else f['content']
            })
        print(json.dumps(output, ensure_ascii=False, indent=2))
