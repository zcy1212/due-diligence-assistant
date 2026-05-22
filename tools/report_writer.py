import os
import json
from pathlib import Path
from typing import List, Dict, Optional

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class ReportWriter:
    def __init__(self, output_dir: str, template_path: Optional[str] = None):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        self.template_path = Path(template_path) if template_path else None

    def set_font(self, run, font_name: str = "宋体", size: int = 12, bold: bool = False):
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.size = Pt(size)
        run.font.bold = bold

    def create_report(self, company_info: Dict, sections: List[str], project_data: List[Dict], 
                      remarks: str = "", filename: str = "尽调报告.docx") -> str:
        if not DOCX_AVAILABLE:
            raise ImportError("需要安装 python-docx: pip install python-docx")
        
        doc = Document()
        
        title = doc.add_heading("尽职调查报告", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"目标公司：{company_info.get('公司全称', '[待补充]')}")
        doc.add_paragraph(f"报告日期：{company_info.get('报告日期', '[待补充]')}")
        doc.add_page_break()
        
        if "声明与说明" in sections:
            self.add_declaration(doc)
        
        if "公司基本情况" in sections:
            self.add_company_info(doc, company_info)
        
        if "公司治理与股权结构" in sections:
            self.add_governance(doc, company_info, project_data)
        
        if "主要业务与资产" in sections:
            self.add_business_assets(doc, project_data)
        
        if "合同与债权债务" in sections:
            self.add_contracts(doc, project_data)
        
        if "劳动与人力资源" in sections:
            self.add_labor(doc)
        
        if "税务与财务" in sections:
            self.add_tax_finance(doc)
        
        if "诉讼仲裁与行政处罚" in sections:
            self.add_litigation(doc)
        
        if "合规与其他事项" in sections:
            self.add_compliance(doc)
        
        if "律师结论与建议" in sections:
            self.add_conclusion(doc, project_data)
        
        if remarks:
            doc.add_page_break()
            doc.add_heading("特殊情况备注", level=1)
            p = doc.add_paragraph(remarks)
            self.set_font(p.runs[0])
        
        doc.add_page_break()
        doc.add_heading("附录", level=1)
        doc.add_heading("一、尽职调查资料清单", level=2)
        doc.add_heading("二、律师核查文件目录", level=2)
        
        output_path = self.output_dir / filename
        doc.save(output_path)
        return str(output_path)

    def create_missing_list(self, missing_items: List[Dict], filename: str = "待补充材料清单.docx") -> str:
        if not DOCX_AVAILABLE:
            raise ImportError("需要安装 python-docx: pip install python-docx")
        
        doc = Document()
        
        title = doc.add_heading("待补充材料清单", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph("请目标公司补充以下材料：")
        doc.add_paragraph("")
        
        for idx, item in enumerate(missing_items, 1):
            p = doc.add_paragraph(f"{idx}. {item.get('类别', '')}：{item.get('材料名称', '')}", style='List Number')
            self.set_font(p.runs[0])
            
            if item.get('说明'):
                p = doc.add_paragraph(f"   说明：{item.get('说明')}")
                self.set_font(p.runs[0])
            
            doc.add_paragraph("")
        
        output_path = self.output_dir / filename
        doc.save(output_path)
        return str(output_path)

    def add_declaration(self, doc):
        doc.add_heading("第一章 声明与说明", level=1)
        p = doc.add_paragraph()
        self.set_font(p.add_run("1. 本报告系基于目标公司提供的文件、资料及公开渠道查询结果出具，调查律师未对资料的真实性、完整性和合法性进行实地核查。"))
        p = doc.add_paragraph()
        self.set_font(p.add_run("2. 本报告所载意见仅供委托方为本次交易目的使用，未经律师事务所同意，不得对外公开或作为其他用途。"))
        doc.add_paragraph("")

    def add_company_info(self, doc, company_info: Dict):
        doc.add_heading("第二章 目标公司基本情况", level=1)
        
        doc.add_heading("1. 公司设立与沿革", level=2)
        p = doc.add_paragraph(f"成立日期：{company_info.get('成立日期', '[待补充]')}")
        self.set_font(p.runs[0])
        p = doc.add_paragraph(f"注册资本：{company_info.get('注册资本', '[待补充]')}万元")
        self.set_font(p.runs[0])
        p = doc.add_paragraph(f"公司类型：{company_info.get('公司类型', '[待补充]')}")
        self.set_font(p.runs[0])
        p = doc.add_paragraph(f"统一社会信用代码：{company_info.get('统一社会信用代码', '[待补充]')}")
        self.set_font(p.runs[0])
        
        doc.add_heading("2. 工商登记与存续情况", level=2)
        p = doc.add_paragraph(f"注册地址：{company_info.get('注册地址', '[待补充]')}")
        self.set_font(p.runs[0])
        p = doc.add_paragraph(f"法定代表人：{company_info.get('法定代表人', '[待补充]')}")
        self.set_font(p.runs[0])
        p = doc.add_paragraph(f"经营范围：{company_info.get('经营范围', '[待补充]')}")
        self.set_font(p.runs[0])
        doc.add_paragraph("")

    def add_governance(self, doc, company_info: Dict, project_data: List[Dict]):
        doc.add_heading("第三章 公司治理与股权结构", level=1)
        
        doc.add_heading("1. 股东及出资情况", level=2)
        p = doc.add_paragraph(f"股东及持股比例：{company_info.get('股东及持股比例', '[待补充]')}")
        self.set_font(p.runs[0])
        
        resolutions = [p for p in project_data if p.get('文件类型') == '决议文件']
        if resolutions:
            doc.add_heading("2. 股东会/董事会决议情况", level=2)
            for res in resolutions:
                p = doc.add_paragraph(f"- {res.get('文件名称', '')}：{res.get('决议内容', '')[:100]}...")
                self.set_font(p.runs[0])
        
        doc.add_paragraph("")

    def add_business_assets(self, doc, project_data: List[Dict]):
        doc.add_heading("第四章 主要业务与资产", level=1)
        doc.add_heading("1. 业务范围与经营模式", level=2)
        doc.add_paragraph("[待补充]")
        doc.add_heading("2. 资产状况", level=2)
        doc.add_paragraph("[待补充]")
        doc.add_paragraph("")

    def add_contracts(self, doc, project_data: List[Dict]):
        doc.add_heading("第五章 合同与债权债务", level=1)
        
        contracts = [p for p in project_data if p.get('文件类型') in ['合同文件', '补充协议']]
        
        if contracts:
            doc.add_heading("1. 重大合同", level=2)
            for contract in contracts:
                p = doc.add_paragraph(f"- {contract.get('文件名称', '')}")
                self.set_font(p.runs[0])
                if contract.get('签约时间'):
                    p = doc.add_paragraph(f"  签约时间：{contract.get('签约时间')}")
                    self.set_font(p.runs[0])
                if contract.get('合同编号'):
                    p = doc.add_paragraph(f"  合同编号：{contract.get('合同编号')}")
                    self.set_font(p.runs[0])
                if contract.get('主要内容摘要'):
                    p = doc.add_paragraph(f"  内容摘要：{contract.get('主要内容摘要')}")
                    self.set_font(p.runs[0])
                doc.add_paragraph("")
        
        doc.add_heading("2. 债务情况", level=2)
        doc.add_paragraph("[待补充]")
        doc.add_paragraph("")

    def add_labor(self, doc):
        doc.add_heading("第六章 劳动与人力资源", level=1)
        doc.add_paragraph("[待补充]")
        doc.add_paragraph("")

    def add_tax_finance(self, doc):
        doc.add_heading("第七章 税务与财务", level=1)
        doc.add_paragraph("[待补充]")
        doc.add_paragraph("")

    def add_litigation(self, doc):
        doc.add_heading("第八章 诉讼、仲裁与行政处罚", level=1)
        doc.add_paragraph("[待补充]")
        doc.add_paragraph("")

    def add_compliance(self, doc):
        doc.add_heading("第九章 合规与其他事项", level=1)
        doc.add_paragraph("[待补充]")
        doc.add_paragraph("")

    def add_conclusion(self, doc, project_data: List[Dict]):
        doc.add_heading("第十章 律师结论与建议", level=1)
        doc.add_heading("1. 目标公司法律合规性总体评价", level=2)
        doc.add_paragraph("[待补充]")
        doc.add_heading("2. 本次交易的主要法律风险", level=2)
        doc.add_paragraph("[待补充]")
        doc.add_heading("3. 风险防控与合规建议", level=2)
        doc.add_paragraph("[待补充]")
        
        missing_count = sum(1 for p in project_data if '[待补充]' in str(p.values()))
        if missing_count > 0:
            p = doc.add_paragraph(f"注：本次尽调中发现 {missing_count} 项信息待补充，详见《待补充材料清单》。")
            self.set_font(p.runs[0])


if __name__ == "__main__":
    import sys
    
    output_dir = sys.argv[1] if len(sys.argv) > 1 else "output"
    action = sys.argv[2] if len(sys.argv) > 2 else "report"
    data_json = sys.argv[3] if len(sys.argv) > 3 else None
    template_path = sys.argv[4] if len(sys.argv) > 4 else None
    
    writer = ReportWriter(output_dir, template_path)
    
    if data_json:
        data = json.loads(data_json)
        if action == "report":
            result = writer.create_report(
                data.get('company_info', {}),
                data.get('sections', []),
                data.get('project_data', []),
                data.get('remarks', '')
            )
            print(result)
        elif action == "missing_list":
            result = writer.create_missing_list(data.get('missing_items', []))
            print(result)
